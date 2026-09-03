"""Tests for DOCX Templates (PR D).

DOCX Templates provide styled document output for different audience types.
"""
from __future__ import annotations

from multi_agent_brief.outputs.templates import TEMPLATES


def test_all_expected_templates_exist():
    expected = {"executive_brief", "research_note", "formal_internal_report"}
    assert set(TEMPLATES.keys()) == expected


def test_default_template_works(tmp_path):
    """Default template should work without issues."""
    md_path = tmp_path / "test.md"
    md_path.write_text("# Test\n\nContent.\n", encoding="utf-8")
    docx_path = tmp_path / "output.docx"

    from multi_agent_brief.outputs.ib_docx import convert
    result = convert(md_path, docx_path, template="default")

    assert result.exists()


def test_heading_inline_markdown_is_stripped_from_docx_headings(tmp_path):
    """LLM-styled heading emphasis must not leak literal Markdown into DOCX."""
    md_path = tmp_path / "test.md"
    md_path.write_text(
        (
            "**美国光储市场周报**\n\n"
            "# **一、核心摘要**\n\n"
            "正文内容。\n\n"
            "### **2.1 美国本土制造与产能扩张**\n\n"
            "更多正文。"
        ),
        encoding="utf-8",
    )
    docx_path = tmp_path / "output.docx"

    from docx import Document
    from multi_agent_brief.outputs.ib_docx import convert

    convert(md_path, docx_path, title="美国光储市场周报", template="research_note")
    doc = Document(str(docx_path))

    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

    assert texts.count("美国光储市场周报") == 1
    assert "一、核心摘要" in headings
    assert "2.1 美国本土制造与产能扩张" in headings
    assert all("*" not in heading for heading in headings)
