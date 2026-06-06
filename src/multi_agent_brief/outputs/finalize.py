from __future__ import annotations

import json
import re
import shutil
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any

from multi_agent_brief.agents.draft_cleanup import strip_claim_citations
from multi_agent_brief.outputs.naming import render_output_stem

_SRC_MARKER_RE = re.compile(r"\[src:[^\]]*\]")


@dataclass
class FinalizeResult:
    """Result of the reader-facing delivery finalization step."""

    status: str
    audited_brief: str
    reader_brief: str
    named_reader_brief: str = ""
    reader_docx: str = ""
    named_reader_docx: str = ""
    docx_generation: str = "not_requested"
    stripped_src_marker_count: int = 0

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def finalize_reader_outputs(
    *,
    output_dir: str | Path,
    project_name: str,
    output_formats: list[str] | tuple[str, ...] | None = None,
    output_footer: str = "",
    output_named_outputs: bool = True,
    output_filename_template: str = "",
    output_filename_tokens: dict[str, str] | None = None,
    docx_template: str = "default",
) -> FinalizeResult:
    """Regenerate reader-facing artifacts from internal audited markdown.

    Agent-assisted workflows may rewrite ``output/intermediate/audited_brief.md``
    after the deterministic pipeline has already produced ``output/brief.md``.
    This function is the final delivery gate: it preserves the cited audited
    artifact for auditability, then writes reader-facing Markdown/DOCX outputs as
    deterministic ``strip_claim_citations(audited_brief)`` derivatives.
    """
    out = Path(output_dir)
    intermediate_dir = out / "intermediate"
    audited_path = intermediate_dir / "audited_brief.md"
    if not audited_path.exists():
        raise FileNotFoundError(
            f"Audited brief not found: {audited_path}. "
            "Run prepare/audit first or write output/intermediate/audited_brief.md."
        )

    out.mkdir(parents=True, exist_ok=True)
    intermediate_dir.mkdir(parents=True, exist_ok=True)

    audited_markdown = audited_path.read_text(encoding="utf-8")
    stripped_count = len(_SRC_MARKER_RE.findall(audited_markdown))
    reader_markdown = strip_claim_citations(audited_markdown)

    brief_path = out / "brief.md"
    brief_path.write_text(reader_markdown, encoding="utf-8")

    named_brief_path: Path | None = None
    if output_named_outputs:
        tokens = dict(output_filename_tokens or {})
        tokens.setdefault("project_name", project_name)
        tokens.setdefault("title", project_name)
        named_stem = render_output_stem(output_filename_template, tokens) if output_filename_template else ""
        if named_stem:
            named_brief_path = out / f"{named_stem}.md"
            if named_brief_path != brief_path:
                named_brief_path.write_text(reader_markdown, encoding="utf-8")

    formats = set(output_formats or ["markdown"])
    docx_status = "not_requested"
    docx_path = out / "brief.docx"
    named_docx_path: Path | None = None
    if "docx" in formats:
        # Avoid leaving a stale rendered file that may still contain internal
        # [src:CLAIM_ID] markers when regeneration fails or dependencies are missing.
        if docx_path.exists():
            docx_path.unlink()
        if named_brief_path is not None:
            possible_named_docx = named_brief_path.with_suffix(".docx")
            if possible_named_docx.exists():
                possible_named_docx.unlink()
        try:
            from multi_agent_brief.outputs.ib_docx import convert

            convert(
                brief_path,
                docx_path,
                title=project_name,
                footer=output_footer or None,
                template=docx_template or "default",
            )
            docx_status = "generated"
            if named_brief_path is not None and named_brief_path.stem != "brief":
                named_docx_path = named_brief_path.with_suffix(".docx")
                shutil.copyfile(docx_path, named_docx_path)
        except ImportError:
            docx_status = "skipped_missing_dependency"
        except Exception:
            docx_status = "failed"
            raise

    result = FinalizeResult(
        status="pass",
        audited_brief=str(audited_path),
        reader_brief=str(brief_path),
        named_reader_brief=str(named_brief_path or ""),
        reader_docx=str(docx_path) if docx_path.exists() else "",
        named_reader_docx=str(named_docx_path or ""),
        docx_generation=docx_status,
        stripped_src_marker_count=stripped_count,
    )

    _assert_reader_artifact_clean(brief_path)
    if named_brief_path and named_brief_path.exists():
        _assert_reader_artifact_clean(named_brief_path)
    if docx_path.exists():
        _assert_docx_artifact_clean(docx_path)
    if named_docx_path and named_docx_path.exists():
        _assert_docx_artifact_clean(named_docx_path)

    report_path = intermediate_dir / "finalize_report.json"
    report_path.write_text(json.dumps(result.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
    _update_audit_report_metadata(
        intermediate_dir / "audit_report.json",
        result,
        named_brief_path=named_brief_path,
    )
    return result


def _assert_reader_artifact_clean(path: Path) -> None:
    text = path.read_text(encoding="utf-8")
    if _SRC_MARKER_RE.search(text):
        raise RuntimeError(f"Reader-facing artifact still contains [src:...] marker: {path}")


def _assert_docx_artifact_clean(path: Path) -> None:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return
    document = Document(str(path))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    if _SRC_MARKER_RE.search(text + "\n" + table_text):
        raise RuntimeError(f"Reader-facing DOCX still contains [src:...] marker: {path}")


def _update_audit_report_metadata(
    audit_report_path: Path,
    result: FinalizeResult,
    *,
    named_brief_path: Path | None,
) -> None:
    if not audit_report_path.exists():
        return
    try:
        payload = json.loads(audit_report_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return
    metadata = payload.setdefault("metadata", {})
    metadata["reader_brief_artifact"] = result.reader_brief
    metadata["reader_brief_transform"] = "strip_claim_citations"
    metadata["reader_brief_finalized"] = True
    metadata["reader_brief_stripped_src_marker_count"] = result.stripped_src_marker_count
    metadata["finalize_report_artifact"] = str(audit_report_path.parent / "finalize_report.json")
    metadata["docx_generation"] = result.docx_generation
    if result.reader_docx:
        metadata["rendered_docx_path"] = result.reader_docx
    if named_brief_path:
        metadata["named_reader_brief_artifact"] = str(named_brief_path)
    audit_report_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
